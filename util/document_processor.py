"""
Document Processor Utility
--------------------------
Handles reading PDF, DOCX, and TXT streams and splitting text 
using LangChain's RecursiveCharacterTextSplitter.
"""

from typing import List
from pypdf import PdfReader
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangChainDocument

def extract_text_from_pdf(file_path_or_bytes) -> str:
    reader = PdfReader(file_path_or_bytes)
    extracted_text = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            extracted_text.append(text)
    return "\n".join(extracted_text)

def extract_text_from_docx(file_path_or_bytes) -> str:
    doc = Document(file_path_or_bytes)
    extracted_text = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(extracted_text)

def chunk_documents(text_content: str, source_name: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[LangChainDocument]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", " ", ""]
    )
    chunks = splitter.split_text(text_content)
    
    return [
        LangChainDocument(
            page_content=chunk,
            metadata={"source": source_name, "chunk_index": i}
        ) for i, chunk in enumerate(chunks)
    ]
